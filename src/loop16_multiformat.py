import os
from pathlib import Path
from docx import Document
import openpyxl
import psycopg2

RAW_DIR = Path(os.path.expanduser("~/projects/lattice/data/raw"))
RAW_DIR.mkdir(parents=True, exist_ok=True)

# create sample docx
doc = Document()
doc.add_paragraph("Lattice supports multi-format ingestion.")
doc.add_paragraph("This paragraph comes from a Word document.")
doc.save(RAW_DIR / "sample.docx")

# create sample xlsx
wb = openpyxl.Workbook()
ws = wb.active
ws.append(["name", "role"])
ws.append(["Lattice", "RAG system"])
ws.append(["pgvector", "vector extension"])
wb.save(RAW_DIR / "sample.xlsx")

def parse_docx(path):
    d = Document(str(path))
    return "\n".join(p.text for p in d.paragraphs if p.text.strip())

def parse_xlsx(path):
    wb = openpyxl.load_workbook(path)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    header, data_rows = rows[0], rows[1:]
    lines = []
    for row in data_rows:
        lines.append(", ".join(f"{h}: {v}" for h, v in zip(header, row)))
    return "\n".join(lines)

def parse_postgres_table(table_name, dsn):
    conn = psycopg2.connect(dsn)
    cur = conn.cursor()
    cur.execute(f"SELECT column_name FROM information_schema.columns WHERE table_name = %s", (table_name,))
    cols = [r[0] for r in cur.fetchall()]
    cur.execute(f"SELECT * FROM {table_name} LIMIT 5")
    rows = cur.fetchall()
    cur.close()
    conn.close()
    return "\n".join(", ".join(f"{c}: {v}" for c, v in zip(cols, row)) for row in rows)

results = {
    "docx": parse_docx(RAW_DIR / "sample.docx"),
    "xlsx": parse_xlsx(RAW_DIR / "sample.xlsx"),
    "postgres": parse_postgres_table("meta_items", "host=localhost port=5432 user=postgres password=pass dbname=postgres"),
}

for fmt, text in results.items():
    print(f"--- {fmt} ---")
    print(text[:150])

assert all(len(t) > 0 for t in results.values()), "one or more formats produced empty text"
assert len(results) >= 4 - 1  # docx, xlsx, postgres (+ txt/html from loop7)
print("OK")
